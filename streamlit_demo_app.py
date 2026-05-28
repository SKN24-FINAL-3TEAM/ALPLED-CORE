import json
import shutil
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT_DIR / "output" / "demo"
JSON_DIR = ROOT_DIR / "json_temp" / "demo"
UI_SAMPLE_DOCX = Path(r"C:\Users\Playdata\Downloads\화면설계서_20260529_022012.docx")


DOC_TYPES = {
    "srs": {
        "tab": "사용자 요구사항 정의서(SRS)",
        "title": "사용자 요구사항 정의서",
        "json_name": "srs_agent_output",
        "docx_name": "사용자 요구사항 정의서",
        "default_input": "./data/requirement_sources/서민금융진흥원 AI기반 통합 플랫폼 구축 사업 제안요청서_final.json",
        "button": "사용자 요구사항 정의서 생성 실행",
    },
    "ui": {
        "tab": "화면설계서",
        "title": "화면설계서",
        "json_name": "ui_design_agent_output",
        "docx_name": "화면설계서",
        "default_input": "./data/interface/requirements/사용자 요구사항 정의서 샘플.json",
        "button": "화면설계서 생성 실행",
    },
    "erd": {
        "tab": "ERD설계서",
        "title": "엔티티 관계 모형 설계서",
        "json_name": "erd_agent_output",
        "docx_name": "엔티티 관계 모형 설계서",
        "default_input": "./data/requirements/requirement.json",
        "button": "ERD 설계서 생성 실행",
    },
    "db": {
        "tab": "DB 설계서",
        "title": "데이터베이스 설계서",
        "json_name": "database_design_agent_output",
        "docx_name": "데이터베이스 설계서",
        "default_input": "./data/requirements/requirement.json",
        "button": "DB 설계서 생성 실행",
    },
    "arch": {
        "tab": "아키텍처 설계서",
        "title": "아키텍처 설계서",
        "json_name": "architecture_agent_output",
        "docx_name": "아키텍처 설계서",
        "default_input": "./data/architecture/architecture_requirements.json",
        "button": "아키텍처 설계서 생성 실행",
    },
    "ts": {
        "tab": "통합시험시나리오",
        "title": "통합 시험 시나리오",
        "json_name": "ts_agent_output",
        "docx_name": "통합 시험 시나리오",
        "default_input": "./data/requirements/requirement.json",
        "button": "통합 시험 시나리오 생성 실행",
    },
}


def timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def docx_output_path(title: str, subdir: str | None = None) -> str:
    output_dir = "./output/demo" if subdir is None else f"./output/demo/{subdir}"
    return f"{output_dir}/{title}_{timestamp()}.docx"


def input_priority_note(default_target: str = "아래 기본 경로") -> None:
    st.info(f"파일을 업로드하면 업로드한 파일로 실행합니다. 업로드하지 않으면 {default_target}를 사용합니다.")


def save_uploaded_file(uploaded_file, prefix: str) -> str | None:
    if uploaded_file is None:
        return None
    upload_dir = ROOT_DIR / "json_temp" / "demo_uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_path = upload_dir / f"{prefix}_{timestamp()}_{Path(uploaded_file.name).name}"
    safe_path.write_bytes(uploaded_file.getbuffer())
    return str(safe_path)


def save_uploaded_files(uploaded_files, prefix: str) -> list[str]:
    return [
        saved
        for uploaded_file in uploaded_files or []
        if (saved := save_uploaded_file(uploaded_file, prefix))
    ]


def read_json(path: str) -> Any:
    file_path = ROOT_DIR / path if not Path(path).is_absolute() else Path(path)
    if not file_path.exists():
        return {}
    try:
        return json.loads(file_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def extract_requirements(data: Any) -> list[dict[str, Any]]:
    if isinstance(data, dict) and isinstance(data.get("requirements"), list):
        return data["requirements"]
    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]
    return []


def setup_doc(title: str) -> Document:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(9)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    return doc


def set_cell_bg(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = str(header)
        set_cell_bg(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            if isinstance(value, list):
                value = "\n".join(str(item) for item in value)
            cell.text = "" if value is None else str(value)


def sample_requirements() -> list[dict[str, Any]]:
    data = read_json(DOC_TYPES["srs"]["default_input"])
    requirements = extract_requirements(data)
    if requirements:
        return requirements[:8]
    return [
        {
            "requirement_id": "REQ-001",
            "requirement_name": "사용자 인증",
            "requirement_type": "기능",
            "description": "시스템은 사용자 로그인 및 권한 검증 기능을 제공하여야 한다.",
            "priority": "상",
        },
        {
            "requirement_id": "REQ-002",
            "requirement_name": "대시보드 조회",
            "requirement_type": "기능",
            "description": "시스템은 업무 현황을 대시보드로 조회할 수 있어야 한다.",
            "priority": "중",
        },
    ]


def build_payload(kind: str, input_path: str) -> dict[str, Any]:
    source_data = read_json(input_path)
    requirements = extract_requirements(source_data) or sample_requirements()
    base_reqs = requirements[:6]

    if kind == "srs":
        final_reqs = [
            {
                "requirement_id": req.get("requirement_id", f"REQ-{idx:03d}"),
                "requirement_name": req.get("requirement_name", "요구사항"),
                "requirement_type": req.get("requirement_type", "기능"),
                "description": req.get("description", "업무 요구사항을 충족하여야 한다."),
                "priority": req.get("priority", "중"),
                "status": "신규",
            }
            for idx, req in enumerate(base_reqs, start=1)
        ]
        return {"final_reqs": final_reqs, "review_reqs": [], "source": input_path}

    if kind == "ui":
        screens = [
            {"screen_id": "UI-001", "screen_name": "로그인", "menu_path": "공통/로그인", "main_actions": ["아이디 입력", "비밀번호 입력", "로그인"]},
            {"screen_id": "UI-002", "screen_name": "업무 대시보드", "menu_path": "업무/대시보드", "main_actions": ["현황 조회", "상세 이동", "알림 확인"]},
            {"screen_id": "UI-003", "screen_name": "신청 관리", "menu_path": "업무/신청관리", "main_actions": ["목록 조회", "조건 검색", "상태 변경"]},
        ]
        return {"screens": screens, "requirements": base_reqs, "source": input_path}

    if kind == "erd":
        entities = [
            {"entity_id": "ENT-001", "entity_name": "USER", "description": "사용자 기본 정보", "columns": ["USER_ID(PK)", "LOGIN_ID", "USER_NAME", "ROLE_CODE"]},
            {"entity_id": "ENT-002", "entity_name": "APPLICATION", "description": "신청 접수 정보", "columns": ["APPLICATION_ID(PK)", "USER_ID(FK)", "STATUS", "CREATED_AT"]},
            {"entity_id": "ENT-003", "entity_name": "DOCUMENT", "description": "첨부 문서 정보", "columns": ["DOCUMENT_ID(PK)", "APPLICATION_ID(FK)", "FILE_NAME"]},
            {"entity_id": "ENT-004", "entity_name": "COUNSELING", "description": "상담 이력 정보", "columns": ["COUNSELING_ID(PK)", "APPLICATION_ID(FK)", "CONTENTS"]},
        ]
        return {"entities": entities, "relationships": ["USER 1:N APPLICATION", "APPLICATION 1:N DOCUMENT", "APPLICATION 1:N COUNSELING"], "source": input_path}

    if kind == "db":
        tables = [
            {"table_id": "TB_USER", "table_name": "사용자", "tablespace": "TS_APP_DATA", "columns": "USER_ID, LOGIN_ID, USER_NAME, ROLE_CODE"},
            {"table_id": "TB_APPLICATION", "table_name": "신청", "tablespace": "TS_APP_DATA", "columns": "APPLICATION_ID, USER_ID, STATUS, CREATED_AT"},
            {"table_id": "TB_DOCUMENT", "table_name": "문서", "tablespace": "TS_FILE_DATA", "columns": "DOCUMENT_ID, APPLICATION_ID, FILE_NAME"},
            {"table_id": "TB_COUNSELING", "table_name": "상담", "tablespace": "TS_APP_DATA", "columns": "COUNSELING_ID, APPLICATION_ID, CONTENTS"},
        ]
        return {"database_name": "ALPLED_DB", "dbms": "PostgreSQL", "tables": tables, "source": input_path}

    if kind == "arch":
        infra = read_json("./data/architecture/infra_spec.json")
        components = [
            {"layer": "Client", "component": "사용자 브라우저", "description": "업무 화면 접근 및 문서 다운로드"},
            {"layer": "Web", "component": "Streamlit/API Gateway", "description": "사용자 입력 수집 및 agent 실행 요청"},
            {"layer": "Application", "component": "AI-DLC Agent Services", "description": "SRS, ERD, DB, ARCH, TS 생성 워크플로우"},
            {"layer": "Data", "component": "Qdrant / PostgreSQL", "description": "RAG 벡터 검색 및 업무 데이터 저장"},
            {"layer": "Model", "component": infra.get("llm_model", "SLLM"), "description": "문서 생성 및 검토"},
        ]
        return {"infra_spec": infra, "components": components, "requirements": base_reqs, "source": input_path}

    scenarios = [
        {"scenario_id": "TS-001", "scenario_name": "로그인 후 대시보드 조회", "test_cases": ["정상 로그인", "권한별 메뉴 표시", "대시보드 데이터 조회"]},
        {"scenario_id": "TS-002", "scenario_name": "신청 접수 및 문서 첨부", "test_cases": ["신청서 저장", "첨부파일 업로드", "접수 상태 변경"]},
        {"scenario_id": "TS-003", "scenario_name": "상담 이력 등록", "test_cases": ["상담 내용 저장", "이력 목록 조회", "변경 로그 확인"]},
    ]
    return {"scenarios": scenarios, "requirements": base_reqs, "source": input_path}


def write_docx(kind: str, payload: dict[str, Any], output_path: Path) -> str:
    config = DOC_TYPES[kind]
    doc = setup_doc(config["title"])
    doc.add_paragraph("본 문서는 시연용 데모 앱에서 샘플 입력 데이터를 기반으로 생성한 산출물입니다.")

    if kind == "srs":
        doc.add_heading("1. 요구사항 목록", level=2)
        rows = [
            [req.get("requirement_id"), req.get("requirement_name"), req.get("requirement_type"), req.get("priority"), req.get("description"), req.get("status")]
            for req in payload["final_reqs"]
        ]
        add_table(doc, ["ID", "요구사항명", "유형", "우선순위", "상세 설명", "상태"], rows)
    elif kind == "ui":
        doc.add_heading("1. 화면 목록", level=2)
        rows = [[s["screen_id"], s["screen_name"], s["menu_path"], s["main_actions"]] for s in payload["screens"]]
        add_table(doc, ["화면 ID", "화면명", "메뉴 경로", "주요 처리"], rows)
    elif kind == "erd":
        doc.add_heading("1. 엔티티 정의", level=2)
        rows = [[e["entity_id"], e["entity_name"], e["description"], e["columns"]] for e in payload["entities"]]
        add_table(doc, ["엔티티 ID", "엔티티명", "설명", "주요 컬럼"], rows)
        doc.add_heading("2. 관계 정의", level=2)
        add_table(doc, ["관계"], [[rel] for rel in payload["relationships"]])
    elif kind == "db":
        doc.add_heading("1. 데이터베이스 개요", level=2)
        add_table(doc, ["항목", "내용"], [["Database", payload["database_name"]], ["DBMS", payload["dbms"]]])
        doc.add_heading("2. 테이블 목록", level=2)
        rows = [[t["table_id"], t["table_name"], t["tablespace"], t["columns"]] for t in payload["tables"]]
        add_table(doc, ["테이블 ID", "테이블명", "테이블스페이스", "주요 컬럼"], rows)
    elif kind == "arch":
        doc.add_heading("1. 시스템 아키텍처", level=2)
        rows = [[c["layer"], c["component"], c["description"]] for c in payload["components"]]
        add_table(doc, ["계층", "구성요소", "설명"], rows)
        doc.add_heading("2. 인프라 기준", level=2)
        infra = payload.get("infra_spec", {})
        add_table(doc, ["항목", "내용"], [[k, v] for k, v in list(infra.items())[:8]])
    else:
        doc.add_heading("1. 통합 시험 시나리오", level=2)
        rows = [[s["scenario_id"], s["scenario_name"], s["test_cases"]] for s in payload["scenarios"]]
        add_table(doc, ["시나리오 ID", "시나리오명", "시험 케이스"], rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return str(output_path)


def create_demo_artifacts(
    kind: str,
    input_path: str,
    *,
    output_json_path: str | None = None,
    output_docx_path: str | None = None,
) -> dict[str, str]:
    ts = timestamp()
    config = DOC_TYPES[kind]
    payload = build_payload(kind, input_path)

    JSON_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    json_path = Path(output_json_path) if output_json_path else JSON_DIR / f"{config['json_name']}_{ts}.json"
    docx_path = Path(output_docx_path) if output_docx_path else OUTPUT_DIR / f"{config['docx_name']}_{ts}.docx"
    if not json_path.is_absolute():
        json_path = ROOT_DIR / json_path
    if not docx_path.is_absolute():
        docx_path = ROOT_DIR / docx_path

    json_path.parent.mkdir(parents=True, exist_ok=True)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    if kind == "ui" and UI_SAMPLE_DOCX.exists():
        shutil.copyfile(UI_SAMPLE_DOCX, docx_path)
    else:
        write_docx(kind, payload, docx_path)

    return {
        "json_path": str(json_path),
        "docx_path": str(docx_path),
        "count": str(_result_count(kind, payload)),
    }


def _result_count(kind: str, payload: dict[str, Any]) -> int:
    if kind == "srs":
        return len(payload.get("final_reqs", []))
    if kind == "ui":
        return len(payload.get("screens", []))
    if kind == "erd":
        return len(payload.get("entities", []))
    if kind == "db":
        return len(payload.get("tables", []))
    if kind == "arch":
        return len(payload.get("components", []))
    return len(payload.get("scenarios", []))


def download_file(path: str, label: str) -> None:
    file_path = Path(path)
    if not file_path.exists():
        st.warning(f"{label} 파일을 찾지 못했습니다: {path}")
        return
    st.download_button(
        f"{label} 다운로드",
        data=file_path.read_bytes(),
        file_name=file_path.name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if file_path.suffix.lower() == ".docx"
        else "application/json",
        key=f"download_{label}_{file_path.name}",
    )


def fake_progress(message: str) -> None:
    progress = st.progress(0, text=message)
    steps = [
        "입력 파일 확인 중...",
        "샘플 요구사항 분석 중...",
        "문서 구조 생성 중...",
        "DOCX 템플릿 반영 중...",
        "결과 파일 저장 중...",
    ]
    for idx in range(50):
        progress.progress((idx + 1) * 2, text=steps[min(idx // 10, len(steps) - 1)])
        time.sleep(0.1)
    progress.empty()


def run_fake_generation(
    kind: str,
    input_path: str,
    output_json_path: str,
    output_docx_path: str,
    *,
    spinner_text: str,
) -> dict[str, str]:
    with st.spinner(spinner_text):
        fake_progress(spinner_text)
        return create_demo_artifacts(
            kind,
            input_path,
            output_json_path=output_json_path,
            output_docx_path=output_docx_path,
        )


def show_result(result: dict[str, str], rows: list[tuple[str, str]], *, count_label: str | None = None) -> None:
    if count_label:
        st.write(count_label)
    for key, label in rows:
        value = result.get(key)
        if value:
            st.write(f"{label}: `{value}`")
            download_file(value, label)


def render_srs_tab() -> None:
    st.subheader("사용자 요구사항 정의서 생성/수정")
    mode = st.radio("실행 모드", ["신규 생성", "수정"], horizontal=True, key="demo_srs_mode")
    save_docx = st.checkbox("DOCX 생성", value=True, key="demo_srs_save_docx")
    st.caption("DOCX 생성 시 아래 `출력 DOCX 경로`에 저장됩니다.")
    output_json_path = st.text_input("출력 JSON 경로", value=f"./json_temp/demo/srs_agent_output_{timestamp()}.json", key="demo_srs_output_json")
    output_reqs_path = st.text_input("final_reqs 출력 경로", value="./output/demo/final_reqs.json", key="demo_srs_output_reqs")
    output_review_path = st.text_input("review_reqs 출력 경로", value="./output/demo/review_reqs.json", key="demo_srs_output_review_reqs")
    output_docx_path = st.text_input("출력 DOCX 경로", value=docx_output_path("사용자 요구사항 정의서"), key="demo_srs_output_docx")

    if mode == "신규 생성":
        input_priority_note("아래 RFP/회의록 기본 경로")
        rfp_upload = st.file_uploader("RFP 원문 또는 분석 JSON 업로드", type=["pdf", "docx", "txt", "md", "json"], key="demo_srs_rfp_upload")
        minutes_upload = st.file_uploader("회의록 TXT 업로드", type=["txt"], key="demo_srs_minutes_upload")
        rfp_json_path = st.text_input(
            "RFP 분석 JSON 기본 경로 (업로드 없을 때 사용)",
            value="./data/requirement_sources/서민금융진흥원 AI기반 통합 플랫폼 구축 사업 제안요청서_final.json",
            key="demo_srs_rfp_path",
        )
        st.checkbox("원문 RFP 업로드 시 요구사항 JSON 자동 추출", value=True, key="demo_srs_extract_rfp_if_needed")
        st.text_input("RFP 추출 JSON 저장 경로", value=f"./json_temp/demo/srs_rfp_extracted_{timestamp()}.json", key="demo_srs_rfp_extract_output_path")
        st.text_input("회의록 TXT 기본 경로 (업로드 없을 때 사용)", value="./data/requirement_sources/meeting_minutes/RFP_변경_회의록.txt", key="demo_srs_minutes_path")

        with st.expander("회의록 먼저 생성", expanded=False):
            input_priority_note("아래 회의록 생성용 RFP 기본 경로")
            st.file_uploader("회의록 생성용 RFP 원문 업로드", type=["pdf", "docx", "txt", "md", "json"], key="demo_srs_minutes_rfp_upload")
            st.text_input("회의록 생성용 RFP 원문 기본 경로 (업로드 없을 때 사용)", value="./data/requirement_sources/RFP/서민금융진흥원 AI기반 통합 플랫폼 구축 사업 제안요청서.pdf", key="demo_srs_minutes_rfp_path")
            st.radio("회의록 종류", ["변경 회의록", "착수 회의록"], horizontal=True, key="demo_srs_minutes_type")
            st.selectbox("생성 모델", ["OpenAI API", "현재 .env LLM"], key="demo_srs_minutes_provider")
            st.text_input("OpenAI 모델명", value="gpt-4.1-mini", key="demo_srs_minutes_openai_model")
            st.number_input("RFP 입력 최대 글자 수", min_value=1000, max_value=30000, value=8000, step=1000, key="demo_srs_minutes_max_chars")
            minutes_output_path = st.text_input("생성 회의록 저장 경로", value="./data/requirement_sources/meeting_minutes/RFP_변경_회의록.txt", key="demo_srs_minutes_output_path")
            if st.button("회의록 생성 실행", key="demo_run_srs_minutes"):
                with st.spinner("회의록 생성 중입니다..."):
                    fake_progress("회의록 생성 중입니다...")
                    path = ROOT_DIR / minutes_output_path
                    path.parent.mkdir(parents=True, exist_ok=True)
                    path.write_text("RFP 변경 회의록\n- 로그인 및 대시보드 요구사항을 보완한다.\n- 문서 산출물 다운로드 기능을 확인한다.", encoding="utf-8")
                st.success("회의록 생성 완료")
                st.write(f"회의록 TXT: `{minutes_output_path}`")
                download_file(str(path), "회의록 TXT")

        if st.button("사용자 요구사항 정의서 신규 생성 실행", type="primary", key="demo_run_srs_generate"):
            uploaded_rfp_path = save_uploaded_file(rfp_upload, "srs_rfp")
            save_uploaded_file(minutes_upload, "srs_minutes")
            result = run_fake_generation(
                "srs",
                uploaded_rfp_path or rfp_json_path,
                output_json_path,
                output_docx_path,
                spinner_text="SRS 신규 생성 중입니다...",
            )
            Path(output_reqs_path).parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(result["json_path"], output_reqs_path)
            Path(output_review_path).write_text("[]", encoding="utf-8")
            st.session_state["demo_result_srs"] = result
            st.success("사용자 요구사항 정의서 신규 생성 완료")
            show_result(
                result,
                [("json_path", "SRS JSON"), ("docx_path", "사용자 요구사항 정의서 DOCX")],
                count_label=f"생성 요구사항: `{result['count']}`건 / 검토 필요: `0`건",
            )
            st.write(f"final_reqs: `{output_reqs_path}`")
            st.write(f"review_reqs: `{output_review_path}`")
    else:
        input_priority_note("아래 기존 SRS 요구사항 기본 경로")
        existing_upload = st.file_uploader("기존 SRS 요구사항 JSON 업로드", type=["json"], key="demo_srs_existing_upload")
        existing_reqs_path = st.text_input("기존 SRS 요구사항 JSON 기본 경로 (업로드 없을 때 사용)", value="./output/final_reqs.json", key="demo_srs_existing_path")
        st.text_area("수정 지시", height=160, key="demo_srs_instruction")
        st.file_uploader("수정 지시 TXT 업로드", type=["txt"], key="demo_srs_instruction_file")
        if st.button("사용자 요구사항 정의서 수정 실행", type="primary", key="demo_run_srs_modify"):
            uploaded_existing_path = save_uploaded_file(existing_upload, "srs_existing")
            result = run_fake_generation("srs", uploaded_existing_path or existing_reqs_path, output_json_path, output_docx_path, spinner_text="SRS 수정 중입니다...")
            st.success("사용자 요구사항 정의서 수정 완료")
            show_result(result, [("json_path", "SRS JSON"), ("docx_path", "사용자 요구사항 정의서 DOCX")], count_label=f"수정 요구사항: `{result['count']}`건 / 검토 필요: `0`건")


def render_ui_design_tab() -> None:
    st.subheader("사용자 인터페이스 설계서 생성")
    input_priority_note("아래 사용자 요구사항/프로토타입 기본 경로 또는 폴더")
    requirement_uploads = st.file_uploader("사용자 요구사항 JSON 업로드", type=["json"], accept_multiple_files=True, key="demo_ui_req_uploads")
    image_uploads = st.file_uploader("프로토타입 이미지 업로드", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="demo_ui_image_uploads")
    requirement_path = st.text_input("사용자 요구사항 JSON 기본 경로/폴더 (업로드 없을 때 사용)", value="./data/interface/requirements", key="demo_ui_req_path")
    st.text_input("프로토타입 이미지 기본 경로/폴더 (업로드 없을 때 사용)", value="./data/interface/prototypes", key="demo_ui_image_path")
    st.number_input("처리할 이미지 수", min_value=1, max_value=50, value=1, step=1, key="demo_ui_max_images")
    st.checkbox("전체 이미지 처리", value=False, key="demo_ui_all_images")
    st.text_input("작업 JSON/모델 원문 저장 폴더", value="./json_temp/interface", key="demo_ui_work_dir")
    output_json_path = st.text_input("출력 통합 JSON 경로", value=f"./json_temp/demo/interface/ui_design_integrated_{timestamp()}.json", key="demo_ui_output_json")
    output_docx_path = st.text_input("출력 DOCX 경로", value=docx_output_path("사용자 인터페이스 설계서", "interface"), key="demo_ui_output_docx")

    if st.button("사용자 인터페이스 설계서 생성 실행", type="primary", key="demo_run_ui_design"):
        uploaded_req_paths = save_uploaded_files(requirement_uploads, "ui_requirement")
        save_uploaded_files(image_uploads, "ui_prototype")
        result = run_fake_generation("ui", uploaded_req_paths[0] if uploaded_req_paths else requirement_path, output_json_path, output_docx_path, spinner_text="사용자 인터페이스 설계서 생성 중입니다...")
        st.success("사용자 인터페이스 설계서 생성 완료")
        st.write(f"화면 수: `{result['count']}`")
        show_result(result, [("json_path", "사용자 인터페이스 설계서 JSON"), ("docx_path", "사용자 인터페이스 설계서 DOCX")])


def render_erd_tab() -> None:
    st.subheader("엔티티 관계 모형 설계서 생성")
    input_priority_note("아래 요구사항 JSON 기본 경로")
    requirement_upload = st.file_uploader("요구사항 JSON 업로드", type=["json"], key="demo_erd_req_upload")
    requirement_path = st.text_input("요구사항 JSON 기본 경로 (업로드 없을 때 사용)", value="./data/requirements/requirement.json", key="demo_erd_req_path")
    col1, col2 = st.columns(2)
    with col1:
        st.checkbox("LLM 사용", value=True, key="demo_erd_use_llm")
    with col2:
        st.checkbox("Mermaid 포함", value=True, key="demo_erd_use_mermaid")
    output_json_path = st.text_input("출력 JSON 경로", value=f"./json_temp/demo/erd_agent_output_{timestamp()}.json", key="demo_erd_output_json")
    output_docx_path = st.text_input("출력 DOCX 경로", value=docx_output_path("엔티티 관계 모형 설계서"), key="demo_erd_output_docx")
    if st.button("엔티티 관계 모형 설계서 생성 실행", type="primary", key="demo_run_erd"):
        uploaded_path = save_uploaded_file(requirement_upload, "erd_requirement")
        result = run_fake_generation("erd", uploaded_path or requirement_path, output_json_path, output_docx_path, spinner_text="ERD 생성 중입니다...")
        st.success("엔티티 관계 모형 설계서 생성 완료")
        show_result(result, [("json_path", "엔티티 관계 모형 JSON"), ("docx_path", "엔티티 관계 모형 설계서 DOCX")])


def render_db_tab() -> None:
    st.subheader("데이터베이스 설계서 생성")
    input_priority_note("아래 요구사항/ERD 기본 경로")
    requirement_upload = st.file_uploader("요구사항 JSON 업로드", type=["json"], key="demo_db_req_upload")
    erd_upload = st.file_uploader("엔티티 관계 모형 설계서 DOCX 업로드", type=["docx"], key="demo_db_erd_upload")
    requirement_path = st.text_input("요구사항 JSON 기본 경로 (업로드 없을 때 사용)", value="./data/requirements/requirement.json", key="demo_db_req_path")
    st.text_input("엔티티 관계 모형 설계서 DOCX 기본 경로 (업로드 없을 때 사용)", value="./output/엔티티 관계 모형 설계서.docx", key="demo_db_erd_path")
    st.checkbox("RAG 보강 사용", value=True, key="demo_db_use_rag")
    output_json_path = st.text_input("출력 JSON 경로", value=f"./json_temp/demo/db_design_output_{timestamp()}.json", key="demo_db_output_json")
    output_docx_path = st.text_input("출력 DOCX 경로", value=docx_output_path("데이터베이스 설계서"), key="demo_db_output_docx")
    if st.button("데이터베이스 설계서 생성 실행", type="primary", key="demo_run_db"):
        uploaded_req_path = save_uploaded_file(requirement_upload, "db_requirement")
        save_uploaded_file(erd_upload, "db_erd")
        result = run_fake_generation("db", uploaded_req_path or requirement_path, output_json_path, output_docx_path, spinner_text="데이터베이스 설계서 생성 중입니다...")
        st.success("데이터베이스 설계서 생성 완료")
        show_result(result, [("json_path", "데이터베이스 설계 JSON"), ("docx_path", "데이터베이스 설계서 DOCX")])


def render_arch_tab() -> None:
    st.subheader("아키텍처 설계서 생성")
    input_priority_note("아래 요구사항/인프라 스펙 기본 경로")
    requirement_upload = st.file_uploader("요구사항 JSON 업로드", type=["json"], key="demo_arch_req_upload")
    infra_upload = st.file_uploader("인프라 스펙 JSON 업로드", type=["json"], key="demo_arch_infra_upload")
    requirement_path = st.text_input("요구사항 JSON 기본 경로 (업로드 없을 때 사용)", value="./data/architecture/architecture_requirements.json", key="demo_arch_req_path")
    st.text_input("인프라 스펙 JSON 기본 경로 (업로드 없을 때 사용)", value="./data/architecture/infra_spec.json", key="demo_arch_infra_path")
    st.checkbox("Mermaid PNG 생성", value=True, key="demo_arch_render_image")
    output_json_path = st.text_input("출력 JSON 경로", value=f"./json_temp/demo/architecture_agent_output_{timestamp()}.json", key="demo_arch_output_json")
    output_docx_path = st.text_input("출력 DOCX 경로", value=docx_output_path("아키텍처 설계서"), key="demo_arch_output_docx")
    output_image_path = st.text_input("출력 이미지 경로", value=f"./output/demo/architecture_diagram_{timestamp()}.png", key="demo_arch_output_image")
    if st.button("아키텍처 설계서 생성 실행", type="primary", key="demo_run_arch"):
        uploaded_req_path = save_uploaded_file(requirement_upload, "arch_requirement")
        save_uploaded_file(infra_upload, "arch_infra")
        result = run_fake_generation("arch", uploaded_req_path or requirement_path, output_json_path, output_docx_path, spinner_text="아키텍처 설계서 생성 중입니다...")
        image_path = ROOT_DIR / output_image_path
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(b"")
        result["image_path"] = str(image_path)
        st.success("아키텍처 설계서 생성 완료")
        show_result(result, [("json_path", "아키텍처 JSON"), ("docx_path", "아키텍처 설계서 DOCX"), ("image_path", "아키텍처 이미지")])


def render_ts_tab() -> None:
    st.subheader("통합 시험 시나리오 생성")
    input_priority_note("아래 요구사항/UI 설계서 기본 경로")
    requirement_upload = st.file_uploader("요구사항 JSON 업로드", type=["json"], key="demo_ts_req_upload")
    ui_uploads = st.file_uploader("UI 설계서 JSON 업로드", type=["json"], accept_multiple_files=True, key="demo_ts_ui_uploads")
    requirement_path = st.text_input("요구사항 JSON 기본 경로 (업로드 없을 때 사용)", value="./data/requirements/requirement.json", key="demo_ts_req_path")
    st.text_area("UI 설계서 JSON 기본 경로 목록 (업로드 없을 때 추가 사용)", value="", height=90, help="여러 개면 줄바꿈으로 입력합니다.", key="demo_ts_ui_paths")
    st.number_input("요구사항별 재시도 횟수", min_value=0, max_value=5, value=1, step=1, key="demo_ts_max_retries")
    output_json_path = st.text_input("출력 JSON 경로", value=f"./json_temp/demo/ts_agent_output_{timestamp()}.json", key="demo_ts_output_json")
    output_docx_path = st.text_input("출력 DOCX 경로", value=docx_output_path("통합 시험 시나리오"), key="demo_ts_output_docx")
    if st.button("통합 시험 시나리오 생성 실행", type="primary", key="demo_run_ts"):
        uploaded_req_path = save_uploaded_file(requirement_upload, "ts_requirement")
        save_uploaded_files(ui_uploads, "ts_ui")
        result = run_fake_generation("ts", uploaded_req_path or requirement_path, output_json_path, output_docx_path, spinner_text="통합 시험 시나리오 생성 중입니다...")
        st.success("통합 시험 시나리오 생성 완료")
        st.write("요약:", {"scenarios": int(result["count"]), "cases": 9})
        show_result(result, [("json_path", "통합 시험 시나리오 JSON"), ("docx_path", "통합 시험 시나리오 DOCX")])


def render_demo_tab(kind: str) -> None:
    config = DOC_TYPES[kind]
    st.subheader(config["title"])
    st.info("시연용 화면입니다. 실제 agent 대신 샘플 입력을 기준으로 5초 후 결과 DOCX/JSON을 생성합니다.")

    with st.form(f"form_{kind}"):
        uploaded = st.file_uploader("입력 파일 업로드(선택)", type=["json", "txt", "pdf", "docx"], key=f"upload_{kind}")
        input_path = st.text_input("기본 입력 경로", value=config["default_input"], key=f"input_{kind}")
        operator = st.text_input("담당자", value="시연 담당자", key=f"operator_{kind}")
        submitted = st.form_submit_button(config["button"], type="primary")

    if submitted:
        if uploaded is not None:
            upload_dir = ROOT_DIR / "json_temp" / "demo_uploads"
            upload_dir.mkdir(parents=True, exist_ok=True)
            safe_path = upload_dir / f"{kind}_{timestamp()}_{Path(uploaded.name).name}"
            safe_path.write_bytes(uploaded.getbuffer())
            input_path = str(safe_path)

        with st.spinner(f"{config['title']} 생성 중입니다..."):
            fake_progress(f"{config['title']} 생성 준비 중...")
            result = create_demo_artifacts(kind, input_path)
            result["operator"] = operator
            st.session_state[f"demo_result_{kind}"] = result

    result = st.session_state.get(f"demo_result_{kind}")
    if result:
        st.success(f"{config['title']} 생성 완료")
        col1, col2, col3 = st.columns(3)
        col1.metric("생성 항목", f"{result['count']}건")
        col2.metric("처리 상태", "완료")
        col3.metric("검토 필요", "0건")

        st.write(f"출력 DOCX 경로: `{result['docx_path']}`")
        st.write(f"출력 JSON 경로: `{result['json_path']}`")
        download_file(result["docx_path"], config["title"])
        download_file(result["json_path"], f"{config['title']} JSON")


def main() -> None:
    st.set_page_config(page_title="ALPLED 산출물 Agent Runner", layout="wide")
    st.title("ALPLED 산출물 Agent Runner")

    tabs = st.tabs([config["tab"] for config in DOC_TYPES.values()])
    renderers = [
        render_srs_tab,
        render_ui_design_tab,
        render_erd_tab,
        render_db_tab,
        render_arch_tab,
        render_ts_tab,
    ]
    for tab, renderer in zip(tabs, renderers):
        with tab:
            renderer()


if __name__ == "__main__":
    main()
